#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
优化版纯本地文档处理器：从PDF/DOCX/TXT源文件中读取并切分原始题目。
生成结构化的中间JSON文件，不包含任何API调用。

本次优化的主要改动点：
1. 彻底移除所有API相关残留代码，确保100%本地专注
2. 增强PDF处理，实现智能调度器：优先pdfplumber，失败时降级OCR（PaddleOCR/easyocr）
3. 优化布局分析函数参数，提升复杂排版题目切分准确率
4. 重构函数结构，提高模块化和可维护性
5. 集中管理配置常量，优化日志系统
6. 增强题目结构化，改进选项解析和题型识别
7. 完善错误处理和资源管理
8. 确保输出JSON格式完全符合规范
9. 内存管理优化：支持大文件分批处理，内存使用监控和限制
10. OCR线程安全：每个线程独立OCR实例，避免资源竞争
11. 断点续传支持：进度跟踪，支持中断后恢复处理
12. 错误处理增强：针对特定异常类型提供精确错误信息
13. 命令行参数支持：完整的argparse集成，支持所有配置
14. 性能监控和统计：处理时间、内存峰值等详细统计
15. 代码质量提升：修复语法错误，优化文件搜索逻辑

输入：通过环境变量 INPUT_DIR 或默认路径（~/Desktop/aaa原始资料）指定的目录。
输出：./output_jsons/extracted_questions.json

JSON格式：
[
  {
    "source_id": "文件名_页码_题目序号",
    "source_file": "原始文件名.pdf",
    "page": 1,
    "raw_text": "完整题目文本，包含题干、选项、表格等所有内容。公式如 $E=mc^2$ 会被保留。",
    "metadata": {
      "extracted_options": ["A. 选项一文本", "B. 选项二文本"],
      "has_table": true,
      "is_scanned": false,
      "module_type": "data_analysis",
      "extraction_method": "pdf_plumber"  // 或 "ocr_paddle", "docx" 等
    }
  }
]
"""

import json
import re
import sys
import tempfile
import time
import hashlib
import errno
import contextlib
import os
import random
import threading
import concurrent.futures
import logging
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple, Union
from datetime import datetime
from io import BytesIO
from dataclasses import dataclass, field
from enum import Enum

# 添加 vendor 目录到路径，以便使用本地依赖
_vendor = Path(__file__).resolve().parent / "vendor"
if _vendor.exists():
    sys.path.insert(0, str(_vendor))

# ---------------------------------------------------------------------------
# 内存监控工具
# ---------------------------------------------------------------------------

class MemoryMonitor:
    """内存使用监控器"""
    
    def __init__(self):
        self._psutil_available = False
        try:
            import psutil  # type: ignore
            self._psutil_available = True
            self.psutil = psutil
        except ImportError:
            # 使用标准logging，因为logger可能还未初始化
            import logging
            logging.getLogger(__name__).debug("psutil不可用，使用简单内存监控")
    
    def get_memory_percent(self) -> float:
        """获取当前内存使用百分比"""
        if self._psutil_available:
            try:
                return self.psutil.virtual_memory().percent
            except Exception:
                return 0.0
        return 0.0
    
    def check_memory_safe(self, max_percent: float = 80.0) -> bool:
        """检查内存使用是否安全（低于指定百分比）"""
        current_percent = self.get_memory_percent()
        if current_percent > max_percent:
            # 使用标准logging
            import logging
            logging.getLogger(__name__).warning(f"内存使用率过高: {current_percent:.1f}% > {max_percent}%")
            return False
        return True
    
    def get_memory_info(self) -> Dict[str, Any]:
        """获取详细内存信息"""
        if self._psutil_available:
            try:
                mem = self.psutil.virtual_memory()
                return {
                    "total_gb": mem.total / (1024**3),
                    "available_gb": mem.available / (1024**3),
                    "used_gb": mem.used / (1024**3),
                    "percent": mem.percent,
                    "free_gb": mem.free / (1024**3)
                }
            except Exception as e:
                # 使用标准logging
                import logging
                logging.getLogger(__name__).debug(f"获取内存信息失败: {str(e)}")
        
        return {"percent": 0.0}

# 全局内存监控器
memory_monitor = MemoryMonitor()

# ---------------------------------------------------------------------------
# 配置管理（集中化）
# ---------------------------------------------------------------------------

@dataclass
class Config:
    """集中管理所有配置常量"""
    # 输入输出
    INPUT_DIR: Path = Path(os.getenv("INPUT_DIR", str(Path.home() / "Desktop" / "aaa原始资料")))
    OUTPUT_DIR: Path = Path(__file__).resolve().parent / "output_jsons"
    LOG_DIR: Path = OUTPUT_DIR / "logs"
    
    # 文本编码
    TEXT_ENCODINGS: Tuple[str, ...] = ("utf-8", "utf-8-sig", "gbk", "gb2312")
    
    # PDF 处理
    PDF_STRICT_MODE: bool = os.getenv("PDF_STRICT_MODE", "0").strip() == "1"
    PDF_MIN_TEXT_LEN: int = int(os.getenv("PDF_MIN_TEXT_LEN", "200").strip() or "200")
    PDF_MAX_Q_DIFF_RATIO: float = float(os.getenv("PDF_MAX_Q_DIFF_RATIO", "0.5").strip() or "0.5")
    PDF_EXTRACT_MAX_RETRIES: int = int(os.getenv("PDF_EXTRACT_MAX_RETRIES", "2").strip() or "2")
    
    # 布局分析参数（优化点）
    LAYOUT_GAP_THRESHOLD_MULTIPLIER: float = 1.8  # 行间距阈值乘数
    LAYOUT_LINE_HEIGHT_MULTIPLIER: float = 2.2    # 行高乘数
    LAYOUT_MIN_OPTIONS_COUNT: int = 2             # 最少选项数
    
    # OCR 配置
    OCR_ENABLE: bool = True
    OCR_PREFERRED_ENGINE: str = "paddleocr"  # "paddleocr" 或 "easyocr"
    OCR_FALLBACK_ENGINE: str = "easyocr"
    OCR_MIN_TEXT_LEN_FOR_FALLBACK: int = 100  # 文本长度小于此值时触发OCR回退
    OCR_CACHE_ENABLE: bool = os.getenv("OCR_CACHE_ENABLE", "1").strip() == "1"  # 是否启用OCR结果缓存
    OCR_CACHE_MAX_SIZE: int = int(os.getenv("OCR_CACHE_MAX_SIZE", "1000").strip() or "1000")  # OCR缓存最大条目数
    OCR_SMART_SELECTION: bool = os.getenv("OCR_SMART_SELECTION", "1").strip() == "1"  # 是否启用智能OCR引擎选择（基于图像质量）
    
    # 并行处理配置
    PARALLEL_WORKERS: int = int(os.getenv("PARALLEL_WORKERS", "4").strip() or "4")
    PARALLEL_ENABLE: bool = os.getenv("PARALLEL_ENABLE", "1").strip() == "1"
    
    # 内存管理配置
    BATCH_SIZE: int = int(os.getenv("BATCH_SIZE", "1000").strip() or "1000")  # 大文件分批处理大小
    MAX_MEMORY_PERCENT: float = float(os.getenv("MAX_MEMORY_PERCENT", "80").strip() or "80")  # 最大内存使用百分比
    
    # 断点续传配置
    RESUME_ENABLE: bool = os.getenv("RESUME_ENABLE", "1").strip() == "1"  # 是否启用断点续传
    
    # 日志
    LOG_LEVEL: str = os.getenv("LOG_LEVEL", "INFO").strip().upper()
    LOG_TO_FILE: bool = True
    LOG_TO_CONSOLE: bool = True
    
    # 题目识别
    MIN_QUESTION_LENGTH: int = 15
    MAX_QUESTION_LENGTH: int = 5000
    
    def __post_init__(self):
        # 创建输出目录
        self.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        self.LOG_DIR.mkdir(parents=True, exist_ok=True)
        
        # 验证日志级别
        if self.LOG_LEVEL not in ["DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"]:
            self.LOG_LEVEL = "INFO"
        
        # 验证并行处理配置
        if self.PARALLEL_WORKERS < 1:
            self.PARALLEL_WORKERS = 1
        elif self.PARALLEL_WORKERS > 16:
            self.PARALLEL_WORKERS = 16
            logger.warning("PARALLEL_WORKERS 超过最大值 16，已调整为 16")
        
        # 验证PDF配置
        if self.PDF_MIN_TEXT_LEN < 10:
            self.PDF_MIN_TEXT_LEN = 10
        if self.PDF_MAX_Q_DIFF_RATIO <= 0 or self.PDF_MAX_Q_DIFF_RATIO > 1:
            self.PDF_MAX_Q_DIFF_RATIO = 0.5
        if self.PDF_EXTRACT_MAX_RETRIES < 1:
            self.PDF_EXTRACT_MAX_RETRIES = 1
        
        # 验证题目长度限制
        if self.MIN_QUESTION_LENGTH < 5:
            self.MIN_QUESTION_LENGTH = 5
        if self.MAX_QUESTION_LENGTH < self.MIN_QUESTION_LENGTH:
            self.MAX_QUESTION_LENGTH = self.MIN_QUESTION_LENGTH * 10
        
        # 验证OCR引擎配置
        valid_ocr_engines = ["paddleocr", "easyocr"]
        if self.OCR_PREFERRED_ENGINE not in valid_ocr_engines:
            self.OCR_PREFERRED_ENGINE = "paddleocr"
        if self.OCR_FALLBACK_ENGINE not in valid_ocr_engines:
            self.OCR_FALLBACK_ENGINE = "easyocr"
        
        # 验证内存管理配置
        if self.BATCH_SIZE < 100:
            self.BATCH_SIZE = 100
        elif self.BATCH_SIZE > 10000:
            self.BATCH_SIZE = 10000
        
        if self.MAX_MEMORY_PERCENT < 10 or self.MAX_MEMORY_PERCENT > 95:
            self.MAX_MEMORY_PERCENT = 80
        
        # 验证OCR缓存配置
        if self.OCR_CACHE_MAX_SIZE < 10:
            self.OCR_CACHE_MAX_SIZE = 10
        elif self.OCR_CACHE_MAX_SIZE > 10000:
            self.OCR_CACHE_MAX_SIZE = 10000
        
        # 配置验证完成（日志记录将在Logger初始化后由主流程处理）

# 全局配置实例
config = Config()

# ---------------------------------------------------------------------------
# 日志系统（优化版）
# ---------------------------------------------------------------------------

class Logger:
    """统一的日志管理器"""
    
    def __init__(self):
        self._cache = set()
        self._cache_max_size = 1000
        
        # 配置Python标准logging
        self.logger = logging.getLogger("document_extractor")
        self.logger.setLevel(getattr(logging, config.LOG_LEVEL))
        
        # 控制台处理器
        if config.LOG_TO_CONSOLE:
            console_handler = logging.StreamHandler(sys.stdout)
            console_handler.setLevel(getattr(logging, config.LOG_LEVEL))
            console_format = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
            console_handler.setFormatter(console_format)
            self.logger.addHandler(console_handler)
        
        # 文件处理器
        if config.LOG_TO_FILE:
            log_file = config.LOG_DIR / f"extraction_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
            file_handler = logging.FileHandler(log_file, encoding='utf-8')
            file_handler.setLevel(getattr(logging, config.LOG_LEVEL))
            file_format = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
            file_handler.setFormatter(file_format)
            self.logger.addHandler(file_handler)
    
    def _should_log(self, level: str, message: str, context: Optional[Dict] = None) -> bool:
        """去重高频日志"""
        if level in ("WARNING", "ERROR"):
            ctx_key = ""
            if context:
                relevant = {k: v for k, v in context.items() if k in ("error", "status", "file", "kind")}
                ctx_key = str(sorted(relevant.items()))
            fingerprint = f"{level}:{message}:{ctx_key}"
            if fingerprint in self._cache:
                return False
            self._cache.add(fingerprint)
            if len(self._cache) > self._cache_max_size:
                self._cache.clear()
        return True
    
    def debug(self, message: str, context: Optional[Dict] = None):
        if self._should_log("DEBUG", message, context):
            self.logger.debug(f"{message} {context if context else ''}")
    
    def info(self, message: str, context: Optional[Dict] = None):
        if self._should_log("INFO", message, context):
            self.logger.info(f"{message} {context if context else ''}")
    
    def warning(self, message: str, context: Optional[Dict] = None):
        if self._should_log("WARNING", message, context):
            self.logger.warning(f"{message} {context if context else ''}")
    
    def error(self, message: str, context: Optional[Dict] = None):
        if self._should_log("ERROR", message, context):
            self.logger.error(f"{message} {context if context else ''}")
    
    def critical(self, message: str, context: Optional[Dict] = None):
        if self._should_log("CRITICAL", message, context):
            self.logger.critical(f"{message} {context if context else ''}")

# 全局日志实例
logger = Logger()

# ---------------------------------------------------------------------------
# OCR 单例管理器（资源管理优化）
# ---------------------------------------------------------------------------

class OCRManager:
    """OCR引擎单例管理器，负责资源加载和释放"""
    
    def __init__(self):
        self._instances = {}  # 主线程实例（向后兼容）
        self._lock = threading.Lock()
        self._available_engines = self._detect_available_engines()
        self._thread_local = threading.local()  # 线程局部存储
        self._cache = {}  # OCR结果缓存 {hash: text}
        self._cache_hits = 0
        self._cache_misses = 0
        self._cache_enabled = config.OCR_CACHE_ENABLE  # 是否启用缓存
        self._max_cache_size = config.OCR_CACHE_MAX_SIZE  # 最大缓存条目数
    
    def _detect_available_engines(self) -> Dict[str, bool]:
        """检测可用的OCR引擎"""
        engines = {}
        
        # 检测PaddleOCR
        try:
            from paddleocr import PaddleOCR  # type: ignore
            engines["paddleocr"] = True
        except ImportError:
            engines["paddleocr"] = False
        
        # 检测EasyOCR
        try:
            import easyocr  # type: ignore
            engines["easyocr"] = True
        except ImportError:
            engines["easyocr"] = False
        
        # 检测PPStructure（表格识别）
        try:
            from paddleocr import PPStructure  # type: ignore
            engines["ppstructure"] = True
        except ImportError:
            engines["ppstructure"] = False
        
        # 检测Pix2Text（公式识别）
        try:
            import pix2text  # type: ignore
            engines["pix2text"] = True
        except ImportError:
            engines["pix2text"] = False
        
        return engines
    
    def get_instance(self, engine_name: str, thread_safe: bool = True) -> Any:
        """获取OCR引擎实例（支持线程安全模式）"""
        # 检查引擎可用性
        if not self._available_engines.get(engine_name, False):
            logger.warning(f"OCR引擎不可用: {engine_name}")
            return None
        
        if thread_safe:
            # 线程安全模式：每个线程有独立实例
            if not hasattr(self._thread_local, 'instances'):
                self._thread_local.instances = {}
            
            thread_instances = self._thread_local.instances
            if engine_name in thread_instances:
                return thread_instances[engine_name]
            
            # 创建线程局部实例
            instance = self._create_engine_instance(engine_name)
            if instance is not None:
                thread_instances[engine_name] = instance
                logger.debug(f"OCR引擎已初始化（线程局部）: {engine_name}", {"thread_id": threading.get_ident()})
            return instance
        else:
            # 向后兼容模式：共享实例（主线程）
            with self._lock:
                if engine_name in self._instances:
                    return self._instances[engine_name]
                
                instance = self._create_engine_instance(engine_name)
                if instance is not None:
                    self._instances[engine_name] = instance
                    logger.debug(f"OCR引擎已初始化（共享）: {engine_name}")
                return instance
    
    def _create_engine_instance(self, engine_name: str) -> Any:
        """创建OCR引擎实例（内部方法）"""
        try:
            if engine_name == "paddleocr":
                from paddleocr import PaddleOCR  # type: ignore
                import logging
                logging.getLogger("ppocr").setLevel(logging.ERROR)
                return PaddleOCR(show_log=False, use_angle_cls=True, lang="ch", use_gpu=False)
            
            elif engine_name == "easyocr":
                import easyocr  # type: ignore
                with open(os.devnull, "w") as fnull:
                    with contextlib.redirect_stdout(fnull), contextlib.redirect_stderr(fnull):
                        return easyocr.Reader(["ch_sim", "en"], gpu=False, verbose=False)
            
            elif engine_name == "ppstructure":
                from paddleocr import PPStructure  # type: ignore
                import logging
                logging.getLogger("ppocr").setLevel(logging.ERROR)
                return PPStructure(show_log=False, use_gpu=False)
            
            elif engine_name == "pix2text":
                import pix2text  # type: ignore
                return pix2text.pix2text.Pix2Text()
            
        except Exception as e:
            logger.error(f"OCR引擎初始化失败: {engine_name}", {"error": str(e)})
        
        return None
    
    def _image_hash(self, img) -> str:
        """计算图像哈希值用于缓存"""
        try:
            # 尝试将图像转换为字节
            import hashlib
            import numpy as np  # type: ignore
            
            if hasattr(img, 'tobytes'):
                # PIL Image 或类似对象
                img_bytes = img.tobytes()
            elif isinstance(img, np.ndarray):
                # numpy 数组
                img_bytes = img.tobytes()
            elif hasattr(img, 'getvalue'):
                # BytesIO 或类似对象
                img_bytes = img.getvalue()
            else:
                # 其他类型，尝试直接使用
                img_bytes = str(img).encode('utf-8')
            
            return hashlib.md5(img_bytes).hexdigest()
        except Exception as e:
            logger.debug(f"图像哈希计算失败，使用替代方案", {"error": str(e)})
            # 使用简单哈希作为后备
            return hashlib.md5(str(id(img)).encode('utf-8')).hexdigest()
    
    def ocr_image_with_cache(self, img, engine_name: str, use_cache: bool = True) -> str:
        """带缓存的OCR图像识别"""
        # 检查缓存是否启用
        if not self._cache_enabled or not use_cache:
            # 缓存未启用或明确要求不使用缓存
            return self._ocr_image_direct(img, engine_name)
        
        # 计算图像哈希
        img_hash = self._image_hash(img)
        cache_key = f"{engine_name}:{img_hash}"
        
        # 检查缓存
        with self._lock:
            if cache_key in self._cache:
                self._cache_hits += 1
                logger.debug(f"OCR缓存命中", {"engine": engine_name, "hash": img_hash[:8]})
                return self._cache[cache_key]
        
        # 缓存未命中，执行OCR
        self._cache_misses += 1
        logger.debug(f"OCR缓存未命中，执行识别", {"engine": engine_name, "hash": img_hash[:8]})
        
        # 执行OCR识别
        text = self._ocr_image_direct(img, engine_name)
        
        # 缓存结果（如果文本非空）
        if text and text.strip():
            with self._lock:
                # 检查缓存大小，如果超过限制则清理最旧的条目
                if len(self._cache) >= self._max_cache_size:
                    # 简单策略：随机删除一个条目（可改进为LRU）
                    key_to_remove = next(iter(self._cache))
                    del self._cache[key_to_remove]
                    logger.debug(f"缓存已满，删除条目: {key_to_remove[:20]}...")
                
                self._cache[cache_key] = text
        
        return text
    
    def _ocr_image_direct(self, img, engine_name: str) -> str:
        """直接OCR图像识别（无缓存）"""
        ocr_instance = self.get_instance(engine_name, thread_safe=True)
        if not ocr_instance:
            return ""
        
        try:
            if engine_name == "paddleocr":
                result = ocr_instance.ocr(img, cls=True)
                lines = []
                for line in result or []:
                    for path_res in line or []:
                        if path_res and len(path_res) >= 2:
                            text = path_res[1][0]
                            if text:
                                lines.append(str(text))
                return "\n".join(lines)
            
            elif engine_name == "easyocr":
                result = ocr_instance.readtext(img, detail=0)
                return "\n".join(result) if result else ""
            
            else:
                return ""
        except Exception as e:
            logger.error(f"OCR识别失败: {engine_name}", {"error": str(e)})
            return ""
    
    def get_cache_stats(self) -> Dict[str, Any]:
        """获取缓存统计信息"""
        with self._lock:
            return {
                "cache_enabled": self._cache_enabled,
                "cache_size": len(self._cache),
                "cache_hits": self._cache_hits,
                "cache_misses": self._cache_misses,
                "hit_rate": self._cache_hits / (self._cache_hits + self._cache_misses) 
                           if (self._cache_hits + self._cache_misses) > 0 else 0.0,
                "max_cache_size": self._max_cache_size
            }
    
    def clear_cache(self):
        """清除OCR缓存"""
        with self._lock:
            self._cache.clear()
            self._cache_hits = 0
            self._cache_misses = 0
            logger.info("OCR缓存已清除")
    
    def release_all(self):
        """释放所有OCR引擎资源（主线程和当前线程）"""
        # 释放主线程实例
        with self._lock:
            for name, instance in list(self._instances.items()):
                if instance:
                    try:
                        if hasattr(instance, 'close'):
                            instance.close()
                        elif hasattr(instance, 'cpu'):
                            del instance
                    except Exception as e:
                        logger.warning(f"释放OCR引擎资源失败: {name}", {"error": str(e)})
                del self._instances[name]
        
        # 释放当前线程实例
        self.release_thread()
        
        logger.info("所有OCR引擎资源已释放")
    
    def release_thread(self):
        """释放当前线程的OCR引擎资源"""
        if hasattr(self._thread_local, 'instances'):
            thread_instances = self._thread_local.instances
            for name, instance in list(thread_instances.items()):
                if instance:
                    try:
                        if hasattr(instance, 'close'):
                            instance.close()
                        elif hasattr(instance, 'cpu'):
                            del instance
                    except Exception as e:
                        logger.warning(f"释放线程OCR引擎资源失败: {name}", {"error": str(e)})
                del thread_instances[name]
            logger.debug(f"线程OCR资源已释放", {"thread_id": threading.get_ident()})

# 全局OCR管理器
ocr_manager = OCRManager()

# ---------------------------------------------------------------------------
# 依赖检测
# ---------------------------------------------------------------------------

def check_dependencies() -> Dict[str, bool]:
    """检查并报告依赖库可用性"""
    deps = {}
    
    try:
        import pdfplumber  # type: ignore
        deps["pdfplumber"] = True
    except ImportError:
        deps["pdfplumber"] = False
        logger.warning("未安装 pdfplumber，PDF解析功能受限")
    
    try:
        from docx import Document as DocxDocument  # type: ignore
        deps["docx"] = True
    except ImportError:
        deps["docx"] = False
        logger.warning("未安装 python-docx，DOCX解析功能受限")
    
    try:
        import fitz  # type: ignore # PyMuPDF
        deps["fitz"] = True
    except ImportError:
        deps["fitz"] = False
    
    # OCR依赖已在OCRManager中检测
    
    return deps

# 全局依赖状态
deps = check_dependencies()

# ---------------------------------------------------------------------------
# 进度跟踪（断点续传支持）
# ---------------------------------------------------------------------------

class ProgressTracker:
    """进度跟踪器，支持断点续传"""
    
    def __init__(self, progress_file: Optional[Path] = None):
        if progress_file is None:
            progress_file = config.OUTPUT_DIR / "progress.json"
        self.progress_file = progress_file
        self._lock = threading.Lock()
        self._progress = self._load_progress()
    
    def _load_progress(self) -> Dict[str, Any]:
        """加载进度文件"""
        if not self.progress_file.exists():
            return {
                "version": "1.0",
                "created": datetime.now().isoformat(),
                "files": {},
                "statistics": {
                    "total_files": 0,
                    "completed_files": 0,
                    "total_questions": 0,
                    "completed_questions": 0
                }
            }
        
        try:
            with open(self.progress_file, "r", encoding="utf-8") as f:
                data = json.load(f)
            # 确保必要字段存在
            if "files" not in data:
                data["files"] = {}
            if "statistics" not in data:
                data["statistics"] = {
                    "total_files": 0,
                    "completed_files": 0,
                    "total_questions": 0,
                    "completed_questions": 0
                }
            return data
        except Exception as e:
            logger.error(f"加载进度文件失败: {str(e)}", {"file": str(self.progress_file)})
            return {
                "version": "1.0",
                "created": datetime.now().isoformat(),
                "files": {},
                "statistics": {
                    "total_files": 0,
                    "completed_files": 0,
                    "total_questions": 0,
                    "completed_questions": 0
                }
            }
    
    def save_progress(self):
        """保存进度到文件"""
        with self._lock:
            self._progress["last_update"] = datetime.now().isoformat()
            atomic_write_json(self.progress_file, self._progress)
            logger.debug("进度文件已保存", {"file": str(self.progress_file)})
    
    def get_file_status(self, file_path: Path) -> str:
        """获取文件处理状态"""
        file_key = str(file_path.resolve())
        return self._progress["files"].get(file_key, {}).get("status", "pending")
    
    def mark_file_started(self, file_path: Path):
        """标记文件开始处理"""
        file_key = str(file_path.resolve())
        with self._lock:
            if file_key not in self._progress["files"]:
                self._progress["files"][file_key] = {}
            self._progress["files"][file_key].update({
                "filename": file_path.name,
                "status": "processing",
                "start_time": datetime.now().isoformat(),
                "processed_questions": 0,
                "total_questions": 0
            })
            self.save_progress()
    
    def mark_file_completed(self, file_path: Path, total_questions: int, processed_questions: int):
        """标记文件处理完成"""
        file_key = str(file_path.resolve())
        with self._lock:
            if file_key not in self._progress["files"]:
                self._progress["files"][file_key] = {}
            self._progress["files"][file_key].update({
                "filename": file_path.name,
                "status": "completed",
                "end_time": datetime.now().isoformat(),
                "total_questions": total_questions,
                "processed_questions": processed_questions,
                "success": True
            })
            # 更新统计
            self._progress["statistics"]["completed_files"] = sum(
                1 for f in self._progress["files"].values() 
                if f.get("status") == "completed"
            )
            self._progress["statistics"]["total_questions"] += total_questions
            self._progress["statistics"]["completed_questions"] += processed_questions
            self.save_progress()
    
    def mark_file_failed(self, file_path: Path, error: str):
        """标记文件处理失败"""
        file_key = str(file_path.resolve())
        with self._lock:
            if file_key not in self._progress["files"]:
                self._progress["files"][file_key] = {}
            self._progress["files"][file_key].update({
                "filename": file_path.name,
                "status": "failed",
                "end_time": datetime.now().isoformat(),
                "error": error,
                "success": False
            })
            self.save_progress()
    
    def filter_pending_files(self, files: List[Path]) -> List[Path]:
        """过滤出待处理的文件（跳过已完成的）"""
        pending_files = []
        for file_path in files:
            status = self.get_file_status(file_path)
            if status != "completed":
                pending_files.append(file_path)
            else:
                logger.info(f"跳过已处理的文件: {file_path.name}")
        return pending_files
    
    def get_statistics(self) -> Dict[str, Any]:
        """获取处理统计信息"""
        return self._progress["statistics"].copy()
    
    def get_overall_progress(self) -> Dict[str, Any]:
        """获取整体进度信息"""
        stats = self.get_statistics()
        total_files = len(self._progress["files"])
        completed_files = sum(1 for f in self._progress["files"].values() 
                             if f.get("status") == "completed")
        
        return {
            "total_files": total_files,
            "completed_files": completed_files,
            "file_progress": f"{(completed_files/total_files*100):.1f}%" if total_files > 0 else "0%",
            "total_questions": stats["total_questions"],
            "completed_questions": stats["completed_questions"],
            "question_progress": f"{(stats['completed_questions']/stats['total_questions']*100):.1f}%" 
                                if stats["total_questions"] > 0 else "0%"
        }

# 全局进度跟踪器
progress_tracker = ProgressTracker()

# ---------------------------------------------------------------------------
# 性能监控和统计
# ---------------------------------------------------------------------------

class PerformanceMonitor:
    """性能监控器，记录处理时间、内存使用等统计信息"""
    
    def __init__(self):
        self.start_time = time.time()
        self.end_time = None
        self.file_timings = {}  # 文件处理时间统计
        self.memory_peaks = []  # 内存使用峰值
        self.total_files_processed = 0
        self.total_questions_processed = 0
        self._lock = threading.Lock()
    
    def record_file_start(self, file_path: Path):
        """记录文件开始处理"""
        file_key = str(file_path.resolve())
        with self._lock:
            self.file_timings[file_key] = {
                "filename": file_path.name,
                "start_time": time.time(),
                "end_time": None,
                "duration": None,
                "questions": 0,
                "memory_peak": 0.0
            }
    
    def record_file_end(self, file_path: Path, questions_count: int):
        """记录文件处理完成"""
        file_key = str(file_path.resolve())
        with self._lock:
            if file_key in self.file_timings:
                end_time = time.time()
                self.file_timings[file_key]["end_time"] = end_time
                self.file_timings[file_key]["duration"] = end_time - self.file_timings[file_key]["start_time"]
                self.file_timings[file_key]["questions"] = questions_count
                self.total_files_processed += 1
                self.total_questions_processed += questions_count
                
                # 记录当前内存使用作为峰值
                memory_percent = memory_monitor.get_memory_percent()
                self.file_timings[file_key]["memory_peak"] = memory_percent
                self.memory_peaks.append(memory_percent)
    
    def record_memory_peak(self):
        """记录当前内存使用峰值"""
        memory_percent = memory_monitor.get_memory_percent()
        self.memory_peaks.append(memory_percent)
    
    def get_statistics(self) -> Dict[str, Any]:
        """获取性能统计信息"""
        total_duration = 0.0
        fastest_file = {"duration": float("inf"), "name": ""}
        slowest_file = {"duration": 0.0, "name": ""}
        
        for file_key, timing in self.file_timings.items():
            if timing["duration"] is not None:
                total_duration += timing["duration"]
                if timing["duration"] < fastest_file["duration"]:
                    fastest_file = {"duration": timing["duration"], "name": timing["filename"]}
                if timing["duration"] > slowest_file["duration"]:
                    slowest_file = {"duration": timing["duration"], "name": timing["filename"]}
        
        avg_duration = total_duration / len(self.file_timings) if self.file_timings else 0
        avg_questions_per_file = self.total_questions_processed / self.total_files_processed if self.total_files_processed > 0 else 0
        avg_questions_per_second = self.total_questions_processed / total_duration if total_duration > 0 else 0
        
        memory_peak_max = max(self.memory_peaks) if self.memory_peaks else 0
        memory_peak_avg = sum(self.memory_peaks) / len(self.memory_peaks) if self.memory_peaks else 0
        
        return {
            "total_files": self.total_files_processed,
            "total_questions": self.total_questions_processed,
            "total_duration_seconds": total_duration,
            "total_duration_human": self._format_duration(total_duration),
            "avg_duration_per_file_seconds": avg_duration,
            "avg_duration_per_file_human": self._format_duration(avg_duration),
            "fastest_file": {
                "name": fastest_file["name"],
                "duration_seconds": fastest_file["duration"],
                "duration_human": self._format_duration(fastest_file["duration"])
            },
            "slowest_file": {
                "name": slowest_file["name"],
                "duration_seconds": slowest_file["duration"],
                "duration_human": self._format_duration(slowest_file["duration"])
            },
            "avg_questions_per_file": round(avg_questions_per_file, 2),
            "avg_questions_per_second": round(avg_questions_per_second, 2),
            "memory_peak_max_percent": round(memory_peak_max, 2),
            "memory_peak_avg_percent": round(memory_peak_avg, 2),
            "parallel_enabled": config.PARALLEL_ENABLE,
            "parallel_workers": config.PARALLEL_WORKERS,
            "batch_size": config.BATCH_SIZE,
            "resume_enabled": config.RESUME_ENABLE
        }
    
    def _format_duration(self, seconds: float) -> str:
        """格式化时间持续时间为人类可读格式"""
        if seconds < 60:
            return f"{seconds:.1f}秒"
        elif seconds < 3600:
            minutes = seconds / 60
            return f"{minutes:.1f}分钟"
        else:
            hours = seconds / 3600
            return f"{hours:.2f}小时"
    
    def print_summary(self):
        """打印性能统计摘要"""
        stats = self.get_statistics()
        logger.info("性能统计摘要", stats)

# 全局性能监控器
performance_monitor = PerformanceMonitor()

# ---------------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------------

def atomic_write_json(path: Path, data: Any) -> None:
    """原子写入JSON文件，避免中断损坏"""
    path.parent.mkdir(parents=True, exist_ok=True)
    content = json.dumps(data, ensure_ascii=False, indent=2)
    try:
        fd, tmp = tempfile.mkstemp(dir=path.parent, suffix=".json", prefix=".tmp_")
        try:
            with os.fdopen(fd, "w", encoding="utf-8") as f:
                f.write(content)
            Path(tmp).replace(path)
            logger.debug(f"JSON文件已保存: {path}")
        except Exception as e:
            Path(tmp).unlink(missing_ok=True)
            raise
    except OSError as e:
        if e.errno == errno.ENOSPC:
            logger.critical("磁盘空间不足，请清理后重试", {"file": str(path)})
            raise
        else:
            logger.error("写入文件失败", {"file": str(path), "error": str(e)})
            raise

def clean_text(text: Optional[str]) -> str:
    """清理冗余字符"""
    if text is None or not isinstance(text, str):
        return ""
    text = str(text)
    # 移除页眉页脚
    text = re.sub(r"第\s*\d+\s*页.*?\d+\s*页", "", text)
    text = re.sub(r"本[,，试卷由粉笔用户].*生成", "", text)
    # 清理冗余标点
    text = re.sub(r"[,，]{2,}", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()

def find_files(input_dir: Path) -> List[Path]:
    """查找输入目录中的所有支持文件"""
    all_files = []
    
    # 递归搜索所有支持的文件类型
    for ext in ["*.pdf", "*.docx", "*.txt"]:
        for file in input_dir.rglob(ext):
            all_files.append(file)
    
    # 按文件名排序
    all_files.sort(key=lambda x: x.name)
    
    logger.debug(f"找到 {len(all_files)} 个文件", {"input_dir": str(input_dir)})
    return all_files

# ---------------------------------------------------------------------------
# 表格转换
# ---------------------------------------------------------------------------

def table_to_markdown(rows: List[List[str]]) -> str:
    """将表格行转为Markdown格式"""
    if not rows or not any(rows):
        return ""
    try:
        clean = [[str(c or "").replace("\n", " ").replace("|", "｜").strip() for c in row] for row in rows if row]
        if not clean:
            return ""
        max_cols = max(len(r) for r in clean)
        for i, row in enumerate(clean):
            if len(row) < max_cols:
                clean[i] = row + [""] * (max_cols - len(row))
        header = clean[0]
        md = "| " + " | ".join(header) + " |\n"
        md += "| " + " | ".join(["---"] * len(header)) + " |\n"
        for row in clean[1:]:
            md += "| " + " | ".join(row[:len(header)]) + " |\n"
        return md
    except (TypeError, IndexError):
        return ""

def html_table_to_markdown(html: str) -> str:
    """HTML表格转Markdown"""
    if not html:
        return ""
    rows = []
    for tr in re.findall(r"<tr[^>]*>([\s\S]*?)</tr>", html, flags=re.I):
        cells = []
        for td in re.findall(r"<t[dh][^>]*>([\s\S]*?)</t[dh]>", tr, flags=re.I):
            text = re.sub(r"<[^>]+>", "", td)
            text = re.sub(r"\s+", " ", text).strip()
            cells.append(text)
        if cells:
            rows.append(cells)
    return table_to_markdown(rows)

# ---------------------------------------------------------------------------
# 题目识别与验证
# ---------------------------------------------------------------------------

def is_question_start_line(text: Optional[str], index: int, lines: List[str]) -> bool:
    """判断是否是题目开始行"""
    if text is None or not isinstance(text, str) or len((t := str(text).strip())) < 2:
        return False
    
    # 数字题号：1. 1、1．(1)
    if re.match(r"^\d+[\.．、,，)]", t):
        return True
    if re.match(r"^[（(]\s*\d+\s*[）)]", t):
        return True
    
    # 题目 N
    if re.search(r"^题目?\s*\d+", t):
        return True
    
    # 中文数字：一、二、
    if re.match(r"^[一二三四五六七八九十百千]+[、．\.]", t):
        return True
    if re.match(r"^[（(][一二三四五六七八九十]+[）)]", t):
        return True
    
    # 【题目】【材料】
    if re.match(r"^【(?:题目|材料)】", t):
        return True
    
    return False

def should_end_question(text: Optional[str], current_question: List[str]) -> bool:
    """判断是否应结束当前题目"""
    if not current_question or text is None:
        return False
    t = str(text).strip() if text else ""
    
    # 答案/解析
    if re.search(r"答案[：:]|解析[：:]|【解析】|参考答案", t, re.IGNORECASE):
        return True
    
    # 下一个题目开始
    if is_question_start_line(text, 0, []):
        return True
    
    # 题目长度限制
    if len(current_question) > 25:
        return True
    
    return False

def is_valid_question(text: Optional[str]) -> bool:
    """验证是否为有效题目"""
    if text is None or not isinstance(text, str):
        return False
    t = str(text).strip()
    
    if len(t) < config.MIN_QUESTION_LENGTH or len(t) > config.MAX_QUESTION_LENGTH:
        return False
    
    has_options = re.search(r"[A-Da-d][\.．,，、\)）]", t)
    has_keywords = re.search(r"选项|答案|选择|正确|错误|下列|哪项|因为|所以|言语|逻辑|资料", t)
    
    if not (has_options or has_keywords):
        return False
    
    # 排除目录、章节标题等
    if re.match(r"^目录$|^索引$|^前言$|^第[一二三四五六七八九十]+章|^第[一二三四五六七八九十]+节", t):
        return False
    
    return True

def extract_question_number(text: Optional[str]) -> int:
    """提取题号用于排序"""
    if text is None or not isinstance(text, str):
        return 99999
    m = re.match(r"(\d+)", str(text).strip())
    if m:
        return int(m.group(1))
    cn_map = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}
    for cn, num in cn_map.items():
        if text.strip().startswith(cn):
            return num
    return 99999

# ---------------------------------------------------------------------------
# PDF 处理核心
# ---------------------------------------------------------------------------

class PDFProcessor:
    """PDF处理器，集成多种提取策略和OCR降级"""
    
    def __init__(self):
        self.extraction_method = "unknown"
    
    def extract_with_pdfplumber(self, pdf_path: Path) -> str:
        """使用pdfplumber提取PDF文本（含表格）"""
        if not deps.get("pdfplumber", False):
            raise RuntimeError("pdfplumber不可用")
        
        all_text = []
        try:
            import pdfplumber  # type: ignore
            with pdfplumber.open(pdf_path) as pdf:
                pages = getattr(pdf, "pages", []) or []
                for page in pages:
                    try:
                        text = page.extract_text() or ""
                        # 提取表格
                        for tbl in page.find_tables() or []:
                            try:
                                data = tbl.extract()
                                if data:
                                    text += "\n\n" + table_to_markdown(data) + "\n"
                            except (AttributeError, TypeError):
                                pass
                        if text and text.strip():
                            all_text.append(text.strip())
                    except (AttributeError, TypeError, IndexError):
                        continue
            self.extraction_method = "pdf_plumber"
            return "\n".join(all_text)
        except Exception as e:
            logger.error("pdfplumber提取失败", {"file": pdf_path.name, "error": str(e)})
            raise
    
    def extract_with_ocr(self, pdf_path: Path) -> str:
        """使用OCR引擎提取扫描PDF文本"""
        if not config.OCR_ENABLE:
            raise RuntimeError("OCR功能未启用")
        
        # 转换PDF为图像
        images = self._pdf_to_images(pdf_path)
        if not images:
            raise RuntimeError("PDF转图像失败")
        
        all_text = []
        used_engines = set()
        
        for img in images:
            # 选择OCR引擎
            if config.OCR_SMART_SELECTION:
                # 智能引擎选择：基于图像质量
                selected_engine = self._select_ocr_engine_by_quality(img)
                fallback_engine = config.OCR_FALLBACK_ENGINE
            else:
                # 传统选择：首选引擎 + 回退
                selected_engine = config.OCR_PREFERRED_ENGINE
                fallback_engine = config.OCR_FALLBACK_ENGINE
            
            # 使用选定的引擎进行OCR
            page_text = self._ocr_image(img, selected_engine)
            
            # 如果结果不理想，尝试回退引擎
            if not page_text or len(page_text.strip()) < config.OCR_MIN_TEXT_LEN_FOR_FALLBACK:
                if fallback_engine != selected_engine:
                    logger.debug(f"OCR结果不理想，尝试回退引擎", {
                        "selected": selected_engine,
                        "fallback": fallback_engine,
                        "text_length": len(page_text.strip()) if page_text else 0
                    })
                    page_text = self._ocr_image(img, fallback_engine)
                    used_engines.add(fallback_engine)
                else:
                    used_engines.add(selected_engine)
            else:
                used_engines.add(selected_engine)
            
            if page_text and page_text.strip():
                all_text.append(page_text.strip())
        
        # 记录使用的引擎
        if len(used_engines) == 1:
            self.extraction_method = f"ocr_{next(iter(used_engines))}"
        else:
            self.extraction_method = f"ocr_mixed_{'_'.join(sorted(used_engines))}"
        
        return "\n".join(all_text)
    
    def _pdf_to_images(self, pdf_path: Path) -> List[Any]:
        """PDF转图像"""
        images = []
        try:
            if deps.get("fitz", False):
                import fitz  # type: ignore
                doc = fitz.open(pdf_path.as_posix())
                for i in range(doc.page_count):
                    page = doc.load_page(i)
                    pix = page.get_pixmap()
                    img_data = pix.tobytes()
                    from PIL import Image  # type: ignore
                    import io
                    img = Image.open(io.BytesIO(img_data))
                    images.append(img)
                doc.close()
            else:
                # 尝试pdf2image
                try:
                    from pdf2image import convert_from_path  # type: ignore
                    images = convert_from_path(pdf_path.as_posix())
                except ImportError:
                    logger.warning("未找到PDF转图像库，OCR功能受限")
        except Exception as e:
            logger.error("PDF转图像失败", {"file": pdf_path.name, "error": str(e)})
        return images
    
    def _assess_image_quality(self, img) -> float:
        """评估图像质量，返回0-1之间的分数（越高表示质量越好）"""
        try:
            import numpy as np  # type: ignore
            from PIL import Image  # type: ignore
            
            # 确保图像是PIL Image
            if not isinstance(img, Image.Image):
                if isinstance(img, np.ndarray):
                    img = Image.fromarray(img)
                else:
                    # 无法评估，返回中等质量
                    return 0.5
            
            # 转换为灰度图
            gray_img = img.convert('L')
            np_img = np.array(gray_img)
            
            # 计算图像清晰度（拉普拉斯方差）
            # 如果OpenCV可用，使用更精确的方法
            try:
                import cv2  # type: ignore
                laplacian_var = cv2.Laplacian(np_img, cv2.CV_64F).var()
                # 归一化：经验阈值，方差>100表示清晰，<50表示模糊
                quality = min(1.0, laplacian_var / 200.0)
                return max(0.0, quality)
            except ImportError:
                # 使用简单的对比度评估（灰度级标准差）
                contrast = np.std(np_img)
                quality = min(1.0, contrast / 100.0)
                return max(0.0, quality)
                
        except Exception as e:
            logger.debug(f"图像质量评估失败，使用默认质量", {"error": str(e)})
            return 0.5
    
    def _select_ocr_engine_by_quality(self, img) -> str:
        """根据图像质量选择OCR引擎"""
        if not config.OCR_SMART_SELECTION:
            # 智能选择未启用，返回首选引擎
            return config.OCR_PREFERRED_ENGINE
        
        try:
            quality = self._assess_image_quality(img)
            logger.debug(f"图像质量评估: {quality:.3f}", {"quality": quality})
            
            # 阈值：质量>0.6使用PaddleOCR，否则使用EasyOCR
            if quality > 0.6:
                return "paddleocr"
            else:
                return "easyocr"
        except Exception as e:
            logger.warning("智能OCR引擎选择失败，使用首选引擎", {"error": str(e)})
            return config.OCR_PREFERRED_ENGINE
    
    def _ocr_image(self, img, engine_name: str) -> str:
        """使用指定OCR引擎识别图像（带缓存）"""
        # 使用OCR管理器的缓存功能
        return ocr_manager.ocr_image_with_cache(img, engine_name, use_cache=True)
    
    def extract_text_mixed(self, pdf_path: Path) -> Tuple[str, str]:
        """智能调度器：优先pdfplumber，失败时降级OCR"""
        extracted_text = ""
        method = "unknown"
        
        try:
            # 首先尝试pdfplumber
            extracted_text = self.extract_with_pdfplumber(pdf_path)
            method = self.extraction_method
            
            # 检查提取结果质量
            if len(extracted_text.strip()) < config.PDF_MIN_TEXT_LEN:
                logger.warning(f"PDF文本过短，尝试OCR降级", {
                    "file": pdf_path.name, 
                    "length": len(extracted_text.strip()),
                    "threshold": config.PDF_MIN_TEXT_LEN
                })
                try:
                    ocr_text = self.extract_with_ocr(pdf_path)
                    if len(ocr_text.strip()) > len(extracted_text.strip()):
                        extracted_text = ocr_text
                        method = self.extraction_method
                except Exception as e:
                    logger.warning("OCR降级失败，继续使用原文本", {"file": pdf_path.name, "error": str(e)})
        
        except Exception as e:
            logger.warning(f"pdfplumber提取失败，尝试OCR", {"file": pdf_path.name, "error": str(e)})
            try:
                extracted_text = self.extract_with_ocr(pdf_path)
                method = self.extraction_method
            except Exception as e2:
                logger.error("所有提取方法均失败", {"file": pdf_path.name, "error": str(e2)})
                raise RuntimeError(f"PDF提取失败: {pdf_path.name}")
        
        return extracted_text, method

# 由于篇幅限制，后续函数将在下一部分继续...

# ---------------------------------------------------------------------------
# 布局分析（优化参数版）
# ---------------------------------------------------------------------------

def median(nums: List[float]) -> float:
    """计算中位数"""
    if not nums:
        return 0.0
    nums_sorted = sorted(nums)
    mid = len(nums_sorted) // 2
    if len(nums_sorted) % 2 == 1:
        return nums_sorted[mid]
    return (nums_sorted[mid - 1] + nums_sorted[mid]) / 2.0

def extract_questions_by_layout(pdf_path: Path, pdf_processor: PDFProcessor) -> List[str]:
    """基于布局的题目切分（优化参数版）"""
    try:
        extracted_text, method = pdf_processor.extract_text_mixed(pdf_path)
        if not extracted_text:
            return []
        
        # 使用pdfplumber获取布局信息
        if not deps.get("pdfplumber", False):
            return []
        
        import pdfplumber  # type: ignore
        pages_lines = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                try:
                    words = page.extract_words(keep_blank_chars=False, use_text_flow=True, extra_attrs=[])
                    y_map = {}
                    for w in words:
                        text = w.get("text", "").strip()
                        if not text:
                            continue
                        x0, top, x1, bottom = w.get("x0", 0), w.get("top", 0), w.get("x1", 0), w.get("bottom", 0)
                        y = (top + bottom) / 2.0
                        key = f"{y:.1f}"
                        if key not in y_map:
                            y_map[key] = []
                        y_map[key].append(text)
                    
                    page_height = page.height
                    line_height = 0.0
                    if y_map:
                        ys = sorted(float(k) for k in y_map.keys())
                        if len(ys) > 1:
                            line_height = min(ys[i] - ys[i - 1] for i in range(1, len(ys)) if ys[i] > ys[i - 1])
                    
                    lines = []
                    for key in sorted(y_map.keys(), key=lambda k: float(k)):
                        y = float(key)
                        line = " ".join(y_map[key])
                        if line:
                            lines.append((y, line))
                    
                    pages_lines.append({
                        "lines": lines, 
                        "height": page_height, 
                        "line_height": line_height
                    })
                except Exception:
                    continue
        
        if not pages_lines:
            return []
        
        results = []
        for page in pages_lines:
            lines_with_y = page.get("lines", [])
            if not lines_with_y:
                continue
            
            lines_only = [l for _, l in lines_with_y]
            ys = [y for y, _ in lines_with_y]
            gaps = [ys[i] - ys[i - 1] for i in range(1, len(ys)) if ys[i] > ys[i - 1]]
            
            line_height = float(page.get("line_height", 0.0) or 0.0)
            median_gap = median(gaps) or 0.0
            
            # 使用可配置的阈值参数
            gap_threshold = max(
                median_gap * config.LAYOUT_GAP_THRESHOLD_MULTIPLIER,
                line_height * config.LAYOUT_LINE_HEIGHT_MULTIPLIER
            )
            
            current = []
            opt_count = 0
            last_y = None
            
            for y, line in lines_with_y:
                line_stripped = line.strip()
                if not line_stripped:
                    continue
                
                # 根据行间距判断题目边界
                if last_y is not None and gap_threshold and (y - last_y) > gap_threshold:
                    if current:
                        q = clean_text("\n".join(current))
                        if is_valid_question(q) and opt_count >= config.LAYOUT_MIN_OPTIONS_COUNT:
                            results.append(q)
                    current = []
                    opt_count = 0
                
                # 题目开始行
                if is_question_start_line(line_stripped, 0, lines_only):
                    if current:
                        q = clean_text("\n".join(current))
                        if is_valid_question(q) and opt_count >= config.LAYOUT_MIN_OPTIONS_COUNT:
                            results.append(q)
                    current = [line_stripped]
                    opt_count = 0
                    last_y = y
                    continue
                
                if current:
                    # 识别选项
                    if re.match(r"^[ABCDabcd][\)）.．、)]\s*", line_stripped):
                        opt_count += 1
                    current.append(line_stripped)
                    last_y = y
            
            # 处理最后一个题目
            if current:
                q = clean_text("\n".join(current))
                if is_valid_question(q) and opt_count >= config.LAYOUT_MIN_OPTIONS_COUNT:
                    results.append(q)
        
        return results
    
    except Exception as e:
        logger.error("布局分析失败", {"file": pdf_path.name, "error": str(e)})
        return []

# ---------------------------------------------------------------------------
# 题目提取策略
# ---------------------------------------------------------------------------

def extract_questions_by_line(lines: List[str]) -> List[str]:
    """基于行的题目切分"""
    questions = []
    current = []
    
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        if not line_stripped and not current:
            continue
        
        if is_question_start_line(line_stripped, i, lines):
            if current:
                q = clean_text("\n".join(current))
                if is_valid_question(q):
                    questions.append(q)
            current = [line_stripped] if line_stripped else []
        elif current:
            if should_end_question(line_stripped, current):
                q = clean_text("\n".join(current))
                if is_valid_question(q):
                    questions.append(q)
                current = [line_stripped] if is_question_start_line(line_stripped, 0, []) else []
            else:
                if line_stripped:
                    current.append(line_stripped)
    
    if current:
        q = clean_text("\n".join(current))
        if is_valid_question(q):
            questions.append(q)
    
    return questions

def extract_questions_by_block(text: str) -> List[str]:
    """基于正则块的题目切分"""
    patterns = [
        r"(?:\n|^)(\d+[\.．、,，)）][^\n]*(?:\n(?!\d+[\.．、,，)）])(?![（(][一二三四五六七八九十]+[）)])[^\n]*)*)",
        r"(?:\n|^)([（(][一二三四五六七八九十]+[）)][^\n]*(?:\n(?![（(][一二三四五六七八九十]+[）)])[^\n]*)*)",
        r"(?:\n|^)([一二三四五六七八九十]+[、．\.][^\n]*(?:\n(?![一二三四五六七八九十]+[、．\.])[^\n]*)*)",
        r"(?:\n|^)(【(?:题目|材料)】[^\n]*(?:\n(?!【(?:题目|材料)】)[^\n]*)*)",
    ]
    
    questions = []
    for pat in patterns:
        for m in re.finditer(pat, text, re.MULTILINE):
            block = m.group(1).strip()
            block = clean_text(block)
            if is_valid_question(block):
                questions.append(block)
    
    return questions

# ---------------------------------------------------------------------------
# 选项解析（增强版）
# ---------------------------------------------------------------------------

def parse_options_from_text(text: str) -> List[str]:
    """从文本中解析选项（增强版）"""
    options = []
    lines = text.split('\n')
    
    # 模式1：标准选项 A. 内容 B. 内容
    option_pattern = r'([A-Da-d])[\.．、\)）]\s*(.+?)(?=(?:\s+[A-Da-d][\.．、\)）])|$)'
    for match in re.finditer(option_pattern, text, re.DOTALL):
        label = match.group(1).upper()
        content = match.group(2).strip()
        options.append(f"{label}. {content}")
    
    # 模式2：括号选项 (A) 内容 (B) 内容
    if not options:
        bracket_pattern = r'\(([A-Da-d])\)\s*(.+?)(?=(?:\([A-Da-d]\))|$)'
        for match in re.finditer(bracket_pattern, text, re.DOTALL):
            label = match.group(1).upper()
            content = match.group(2).strip()
            options.append(f"{label}. {content}")
    
    # 模式3：跨行选项
    if not options:
        current_option = None
        current_content = []
        for line in lines:
            line = line.strip()
            m = re.match(r'^([A-Da-d])[\.．、\)）]\s*(.+)$', line)
            if m:
                if current_option is not None:
                    options.append(f"{current_option}. {' '.join(current_content)}")
                current_option = m.group(1).upper()
                current_content = [m.group(2)]
            elif current_option is not None and line:
                current_content.append(line)
        
        if current_option is not None:
            options.append(f"{current_option}. {' '.join(current_content)}")
    
    return options

# ---------------------------------------------------------------------------
# 模块类型识别（增强版）
# ---------------------------------------------------------------------------

def detect_module_type(filename: str, content: str) -> str:
    """识别题目模块类型（基于文件名和内容关键词）"""
    name = Path(filename).stem.lower()
    
    # 基于文件名
    if "资料分析" in name:
        return "data_analysis"
    if "图形推理" in name:
        return "figure_reasoning"
    if "定义判断" in name:
        return "definition_judgement"
    if "数量关系" in name:
        return "quantity_relation"
    if "类比推理" in name:
        return "analogy_reasoning"
    if "语句表达" in name:
        return "sentence_expression"
    if "逻辑判断" in name:
        return "logic_judgement"
    if "逻辑填空" in name:
        return "logic_fill"
    if "阅读理解" in name:
        return "reading_comprehension"
    
    # 基于内容关键词
    content_lower = content.lower()
    if any(keyword in content_lower for keyword in ["根据资料", "同比增长", "环比增长", "增长率", "百分点"]):
        return "data_analysis"
    if any(keyword in content_lower for keyword in ["图形", "图案", "形状", "对称", "旋转"]):
        return "figure_reasoning"
    if any(keyword in content_lower for keyword in ["定义", "概念", "含义", "内涵", "外延"]):
        return "definition_judgement"
    if any(keyword in content_lower for keyword in ["数学", "计算", "方程", "公式", "几何"]):
        return "quantity_relation"
    if any(keyword in content_lower for keyword in ["逻辑", "推理", "前提", "结论", "假设"]):
        return "logic_judgement"
    
    return "unknown"

# ---------------------------------------------------------------------------
# 题目去重与合并
# ---------------------------------------------------------------------------

def question_signature(question: str) -> str:
    """生成题目签名用于去重"""
    if not question or not isinstance(question, str):
        return ""
    s = re.sub(r"\s+", "", str(question))
    s = re.sub(r"[\.．、,，\s]", "", s)
    return s[:120] if s else ""

def deduplicate_questions(questions: List[str]) -> List[str]:
    """题目去重，保留较长版本"""
    by_sig = {}
    for q in questions:
        if not q or not isinstance(q, str):
            continue
        s = question_signature(q)
        if not s:
            continue
        if s not in by_sig or len(q) > len(by_sig[s]):
            by_sig[s] = q
    
    result = list(by_sig.values())
    return sorted(result, key=extract_question_number)

def merge_questions_precise(q1: List[str], q2: List[str], strict: bool = False) -> List[str]:
    """精确合并题目列表"""
    if not strict:
        return q1 + q2
    
    map1 = {question_signature(q): q for q in q1 if question_signature(q)}
    map2 = {question_signature(q): q for q in q2 if question_signature(q)}
    
    common = []
    for sig, q in map1.items():
        if sig in map2:
            q2v = map2[sig]
            common.append(q if len(q) >= len(q2v) else q2v)
    
    if not common and (q1 or q2):
        logger.warning("严格合并后为空，尝试降级为并集", {"q1": len(q1), "q2": len(q2)})
        return list(set(q1 + q2))
    
    return common

# ---------------------------------------------------------------------------
# 文档格式处理
# ---------------------------------------------------------------------------

def extract_from_docx(docx_path: Path) -> str:
    """从DOCX提取文本"""
    if not deps.get("docx", False):
        raise RuntimeError("python-docx不可用")
    
    try:
        from docx import Document  # type: ignore
        doc = Document(docx_path.as_posix())
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # 处理表格
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(cells)
            if rows:
                full_text.append("\n" + table_to_markdown(rows) + "\n")
        
        return "\n".join(full_text)
    except Exception as e:
        logger.error("DOCX提取失败", {"file": docx_path.name, "error": str(e)})
        raise

def load_from_txt(txt_path: Path) -> str:
    """从TXT文件读取文本（多编码支持）"""
    text = ""
    for enc in config.TEXT_ENCODINGS:
        try:
            text = txt_path.read_text(encoding=enc)
            break
        except (UnicodeDecodeError, LookupError, OSError):
            continue
    
    if not text:
        try:
            text = txt_path.read_text(encoding="utf-8", errors="replace")
        except OSError:
            return ""
    
    return text

# ---------------------------------------------------------------------------
# 核心提取函数
# ---------------------------------------------------------------------------

def extract_questions_from_pdf(pdf_path: Path) -> Tuple[List[str], str]:
    """从PDF提取题目，返回题目列表和提取方法"""
    pdf_processor = PDFProcessor()
    
    try:
        # 提取文本
        raw_text, method = pdf_processor.extract_text_mixed(pdf_path)
        if not raw_text:
            logger.warning("PDF文本为空", {"file": pdf_path.name})
            return [], method
        
        # 特殊处理逻辑填空
        if "逻辑填空" in pdf_path.stem:
            raw_text = re.sub(r"[_＿]+", "*****", raw_text)
        
        lines = [ln.strip() for ln in raw_text.split("\n") if ln.strip() and len(ln.strip()) > 2]
        
        # 多策略提取
        q_by_line = extract_questions_by_line(lines)
        q_by_block = extract_questions_by_block(raw_text)
        
        # 布局分析（严格模式时使用）
        q_by_layout = []
        if config.PDF_STRICT_MODE:
            q_by_layout = extract_questions_by_layout(pdf_path, pdf_processor)
        
        # 合并结果
        merged = merge_questions_precise(q_by_line, q_by_block, strict=config.PDF_STRICT_MODE)
        if config.PDF_STRICT_MODE and q_by_layout:
            merged = merge_questions_precise(merged, q_by_layout, strict=True)
        
        # 去重
        unique = deduplicate_questions(merged)
        
        logger.info(f"PDF提取完成", {
            "file": pdf_path.name,
            "method": method,
            "by_line": len(q_by_line),
            "by_block": len(q_by_block),
            "by_layout": len(q_by_layout),
            "final": len(unique)
        })
        
        return unique, method
    
    except Exception as e:
        logger.error("PDF题目提取失败", {"file": pdf_path.name, "error": str(e)})
        return [], "error"

def extract_questions_from_docx(docx_path: Path) -> List[str]:
    """从DOCX提取题目"""
    try:
        raw_text = extract_from_docx(docx_path)
        if not raw_text:
            return []
        
        lines = [ln.strip() for ln in raw_text.split("\n") if ln.strip() and len(ln.strip()) > 2]
        
        q_by_line = extract_questions_by_line(lines)
        q_by_block = extract_questions_by_block(raw_text)
        
        # 简单去重
        seen = set()
        unique = []
        for q in q_by_line + q_by_block:
            if not q or not isinstance(q, str) or len(str(q).strip()) < 15:
                continue
            sig = question_signature(q)
            if sig not in seen:
                seen.add(sig)
                unique.append(q)
        
        logger.info(f"DOCX提取完成", {"file": docx_path.name, "count": len(unique)})
        return sorted(unique, key=extract_question_number)
    
    except Exception as e:
        logger.error("DOCX题目提取失败", {"file": docx_path.name, "error": str(e)})
        return []

def extract_questions_from_txt(txt_path: Path) -> List[str]:
    """从TXT提取题目"""
    try:
        raw_text = load_from_txt(txt_path)
        if not raw_text:
            return []
        
        # 按空行分割
        blocks = [b.strip() for b in re.split(r"\n\s*\n", raw_text) 
                 if b and b.strip() and is_valid_question(b.strip())]
        
        logger.info(f"TXT提取完成", {"file": txt_path.name, "count": len(blocks)})
        return blocks
    
    except Exception as e:
        logger.error("TXT题目提取失败", {"file": txt_path.name, "error": str(e)})
        return []

# ---------------------------------------------------------------------------
# 单文件处理函数（用于并行化）
# ---------------------------------------------------------------------------

def process_single_pdf(pdf_file: Path) -> List[Dict[str, Any]]:
    """处理单个PDF文件并返回结构化题目列表"""
    items = []
    try:
        questions, method = extract_questions_from_pdf(pdf_file)
        
        # 记录内存使用情况
        initial_memory = memory_monitor.get_memory_info()
        logger.debug(f"开始处理PDF文件，初始内存: {initial_memory.get('percent', 0):.1f}%", 
                    {"file": pdf_file.name, "questions": len(questions)})
        
        # 分批处理题目（避免内存峰值）
        batch_size = config.BATCH_SIZE
        total_questions = len(questions)
        
        for batch_start in range(0, total_questions, batch_size):
            batch_end = min(batch_start + batch_size, total_questions)
            batch_questions = questions[batch_start:batch_end]
            
            logger.debug(f"处理题目批次 {batch_start//batch_size + 1}/{(total_questions + batch_size - 1)//batch_size}",
                        {"start": batch_start, "end": batch_end, "size": len(batch_questions)})
            
            # 处理当前批次的题目
            for idx_in_batch, q_text in enumerate(batch_questions):
                idx = batch_start + idx_in_batch
                
                # 解析选项
                options = parse_options_from_text(q_text)
                
                # 检测是否有表格
                has_table = '|' in q_text or '---' in q_text
                
                # 检测是否为扫描件
                is_scanned = method.startswith("ocr_")
                
                # 识别模块类型
                module_type = detect_module_type(pdf_file.name, q_text)
                
                item = {
                    "source_id": f"{pdf_file.stem}_p1_q{idx+1:03d}",
                    "source_file": pdf_file.name,
                    "page": 1,  # 简化处理，实际可扩展为多页
                    "raw_text": q_text,
                    "metadata": {
                        "extracted_options": options,
                        "has_table": has_table,
                        "is_scanned": is_scanned,
                        "module_type": module_type,
                        "extraction_method": method
                    }
                }
                items.append(item)
            
            # 每批处理后检查内存使用
            if not memory_monitor.check_memory_safe(config.MAX_MEMORY_PERCENT):
                logger.warning(f"内存使用超过阈值 {config.MAX_MEMORY_PERCENT}%，继续处理但请注意风险",
                              {"file": pdf_file.name, "batch": batch_start//batch_size + 1})
            
            # 每处理10批或最后一批时记录进度
            if ((batch_start // batch_size + 1) % 10 == 0 or 
                batch_end == total_questions):
                current_memory = memory_monitor.get_memory_info()
                progress_percent = (batch_end / total_questions) * 100
                logger.info(f"PDF处理进度", {
                    "file": pdf_file.name,
                    "progress": f"{progress_percent:.1f}%",
                    "processed": batch_end,
                    "total": total_questions,
                    "memory_percent": f"{current_memory.get('percent', 0):.1f}%"
                })
        
        final_memory = memory_monitor.get_memory_info()
        memory_used = final_memory.get("percent", 0) - initial_memory.get("percent", 0)
        logger.info(f"PDF处理完成", {
            "file": pdf_file.name, 
            "questions": len(questions),
            "memory_used": f"{memory_used:.1f}%",
            "final_memory": f"{final_memory.get('percent', 0):.1f}%"
        })
        return items
    
    except (FileNotFoundError, PermissionError) as e:
        logger.error(f"PDF文件访问失败", {"file": pdf_file.name, "error": str(e)})
        return []
    except (MemoryError, OSError) as e:
        logger.error(f"PDF处理系统错误", {"file": pdf_file.name, "error": str(e)})
        return []
    except ValueError as e:
        logger.error(f"PDF内容解析错误", {"file": pdf_file.name, "error": str(e)})
        return []
    except Exception as e:
        logger.error(f"PDF处理未知错误", {"file": pdf_file.name, "error": str(e)})
        return []

def process_single_docx(docx_file: Path) -> List[Dict[str, Any]]:
    """处理单个DOCX文件并返回结构化题目列表"""
    items = []
    try:
        questions = extract_questions_from_docx(docx_file)
        
        # 记录内存使用情况
        initial_memory = memory_monitor.get_memory_info()
        logger.debug(f"开始处理DOCX文件，初始内存: {initial_memory.get('percent', 0):.1f}%", 
                    {"file": docx_file.name, "questions": len(questions)})
        
        # 分批处理题目（避免内存峰值）
        batch_size = config.BATCH_SIZE
        total_questions = len(questions)
        
        for batch_start in range(0, total_questions, batch_size):
            batch_end = min(batch_start + batch_size, total_questions)
            batch_questions = questions[batch_start:batch_end]
            
            logger.debug(f"处理题目批次 {batch_start//batch_size + 1}/{(total_questions + batch_size - 1)//batch_size}",
                        {"start": batch_start, "end": batch_end, "size": len(batch_questions)})
            
            # 处理当前批次的题目
            for idx_in_batch, q_text in enumerate(batch_questions):
                idx = batch_start + idx_in_batch
                
                options = parse_options_from_text(q_text)
                has_table = '|' in q_text or '---' in q_text
                module_type = detect_module_type(docx_file.name, q_text)
                
                item = {
                    "source_id": f"{docx_file.stem}_q{idx+1:03d}",
                    "source_file": docx_file.name,
                    "page": None,
                    "raw_text": q_text,
                    "metadata": {
                        "extracted_options": options,
                        "has_table": has_table,
                        "is_scanned": False,
                        "module_type": module_type,
                        "extraction_method": "docx"
                    }
                }
                items.append(item)
            
            # 每批处理后检查内存使用
            if not memory_monitor.check_memory_safe(config.MAX_MEMORY_PERCENT):
                logger.warning(f"内存使用超过阈值 {config.MAX_MEMORY_PERCENT}%，继续处理但请注意风险",
                              {"file": docx_file.name, "batch": batch_start//batch_size + 1})
            
            # 每处理10批或最后一批时记录进度
            if ((batch_start // batch_size + 1) % 10 == 0 or 
                batch_end == total_questions):
                current_memory = memory_monitor.get_memory_info()
                progress_percent = (batch_end / total_questions) * 100
                logger.info(f"DOCX处理进度", {
                    "file": docx_file.name,
                    "progress": f"{progress_percent:.1f}%",
                    "processed": batch_end,
                    "total": total_questions,
                    "memory_percent": f"{current_memory.get('percent', 0):.1f}%"
                })
        
        final_memory = memory_monitor.get_memory_info()
        memory_used = final_memory.get("percent", 0) - initial_memory.get("percent", 0)
        logger.info(f"DOCX处理完成", {
            "file": docx_file.name, 
            "questions": len(questions),
            "memory_used": f"{memory_used:.1f}%",
            "final_memory": f"{final_memory.get('percent', 0):.1f}%"
        })
        return items
    
    except (FileNotFoundError, PermissionError) as e:
        logger.error(f"DOCX文件访问失败", {"file": docx_file.name, "error": str(e)})
        return []
    except (MemoryError, OSError) as e:
        logger.error(f"DOCX处理系统错误", {"file": docx_file.name, "error": str(e)})
        return []
    except ValueError as e:
        logger.error(f"DOCX内容解析错误", {"file": docx_file.name, "error": str(e)})
        return []
    except Exception as e:
        logger.error(f"DOCX处理未知错误", {"file": docx_file.name, "error": str(e)})
        return []

def process_single_txt(txt_file: Path) -> List[Dict[str, Any]]:
    """处理单个TXT文件并返回结构化题目列表"""
    items = []
    try:
        questions = extract_questions_from_txt(txt_file)
        
        # 记录内存使用情况
        initial_memory = memory_monitor.get_memory_info()
        logger.debug(f"开始处理TXT文件，初始内存: {initial_memory.get('percent', 0):.1f}%", 
                    {"file": txt_file.name, "questions": len(questions)})
        
        # 分批处理题目（避免内存峰值）
        batch_size = config.BATCH_SIZE
        total_questions = len(questions)
        
        for batch_start in range(0, total_questions, batch_size):
            batch_end = min(batch_start + batch_size, total_questions)
            batch_questions = questions[batch_start:batch_end]
            
            logger.debug(f"处理题目批次 {batch_start//batch_size + 1}/{(total_questions + batch_size - 1)//batch_size}",
                        {"start": batch_start, "end": batch_end, "size": len(batch_questions)})
            
            # 处理当前批次的题目
            for idx_in_batch, q_text in enumerate(batch_questions):
                idx = batch_start + idx_in_batch
                
                options = parse_options_from_text(q_text)
                has_table = '|' in q_text or '---' in q_text
                module_type = detect_module_type(txt_file.name, q_text)
                
                item = {
                    "source_id": f"{txt_file.stem}_q{idx+1:03d}",
                    "source_file": txt_file.name,
                    "page": None,
                    "raw_text": q_text,
                    "metadata": {
                        "extracted_options": options,
                        "has_table": has_table,
                        "is_scanned": False,
                        "module_type": module_type,
                        "extraction_method": "txt"
                    }
                }
                items.append(item)
            
            # 每批处理后检查内存使用
            if not memory_monitor.check_memory_safe(config.MAX_MEMORY_PERCENT):
                logger.warning(f"内存使用超过阈值 {config.MAX_MEMORY_PERCENT}%，继续处理但请注意风险",
                              {"file": txt_file.name, "batch": batch_start//batch_size + 1})
            
            # 每处理10批或最后一批时记录进度
            if ((batch_start // batch_size + 1) % 10 == 0 or 
                batch_end == total_questions):
                current_memory = memory_monitor.get_memory_info()
                progress_percent = (batch_end / total_questions) * 100
                logger.info(f"TXT处理进度", {
                    "file": txt_file.name,
                    "progress": f"{progress_percent:.1f}%",
                    "processed": batch_end,
                    "total": total_questions,
                    "memory_percent": f"{current_memory.get('percent', 0):.1f}%"
                })
        
        final_memory = memory_monitor.get_memory_info()
        memory_used = final_memory.get("percent", 0) - initial_memory.get("percent", 0)
        logger.info(f"TXT处理完成", {
            "file": txt_file.name, 
            "questions": len(questions),
            "memory_used": f"{memory_used:.1f}%",
            "final_memory": f"{final_memory.get('percent', 0):.1f}%"
        })
        return items
    
    except (FileNotFoundError, PermissionError) as e:
        logger.error(f"TXT文件访问失败", {"file": txt_file.name, "error": str(e)})
        return []
    except (MemoryError, OSError) as e:
        logger.error(f"TXT处理系统错误", {"file": txt_file.name, "error": str(e)})
        return []
    except ValueError as e:
        logger.error(f"TXT内容解析错误", {"file": txt_file.name, "error": str(e)})
        return []
    except Exception as e:
        logger.error(f"TXT处理未知错误", {"file": txt_file.name, "error": str(e)})
        return []

def process_files_parallel(files: List[Path], processor_func, file_type: str) -> List[Dict[str, Any]]:
    """并行处理文件列表（支持断点续传）"""
    if not files:
        return []
    
    # 断点续传：过滤已完成的文件
    if config.RESUME_ENABLE:
        pending_files = progress_tracker.filter_pending_files(files)
        if not pending_files:
            logger.info(f"所有{file_type}文件已完成，跳过处理")
            return []
        
        logger.info(f"文件处理状态（断点续传已启用）", {
            "total": len(files),
            "pending": len(pending_files),
            "completed": len(files) - len(pending_files)
        })
    else:
        pending_files = files
        logger.info(f"文件处理状态（断点续传已禁用）", {"total": len(files)})
    
    # 创建处理器（根据配置决定是否跟踪进度）
    if config.RESUME_ENABLE:
        def tracked_processor(file: Path) -> List[Dict[str, Any]]:
            """带进度跟踪的处理器"""
            try:
                # 性能监控：记录开始
                performance_monitor.record_file_start(file)
                
                # 标记文件开始处理
                progress_tracker.mark_file_started(file)
                
                # 调用原始处理器
                items = processor_func(file)
                
                # 标记文件完成
                total_questions = len(items)
                processed_questions = total_questions  # 假设全部成功处理
                progress_tracker.mark_file_completed(file, total_questions, processed_questions)
                
                # 性能监控：记录完成
                performance_monitor.record_file_end(file, total_questions)
                
                return items
                
            except Exception as e:
                # 标记文件失败
                progress_tracker.mark_file_failed(file, str(e))
                raise
    else:
        # 无进度跟踪的处理器（仍需性能监控）
        def tracked_processor(file: Path) -> List[Dict[str, Any]]:
            """无进度跟踪但带性能监控的处理器"""
            try:
                # 性能监控：记录开始
                performance_monitor.record_file_start(file)
                
                # 调用原始处理器
                items = processor_func(file)
                
                # 性能监控：记录完成
                total_questions = len(items)
                performance_monitor.record_file_end(file, total_questions)
                
                return items
                
            except Exception as e:
                raise
    
    if not config.PARALLEL_ENABLE or len(pending_files) < 2:
        # 顺序处理（带进度指示）
        all_items = []
        total_files = len(pending_files)
        logger.info(f"开始顺序处理{file_type}文件", {"file_count": total_files})
        
        for idx, file in enumerate(pending_files, 1):
            try:
                items = tracked_processor(file)
                all_items.extend(items)
                
                # 进度更新
                if idx % 5 == 0 or idx == total_files:
                    progress_percent = (idx / total_files) * 100
                    overall = progress_tracker.get_overall_progress()
                    logger.info(f"{file_type}文件处理进度", {
                        "completed": idx,
                        "total": total_files,
                        "progress": f"{progress_percent:.1f}%",
                        "overall_files": f"{overall['completed_files']}/{overall['total_files']}",
                        "overall_questions": f"{overall['completed_questions']}/{overall['total_questions']}"
                    })
                    
            except Exception as e:
                logger.error(f"{file_type}文件处理失败", {"file": file.name, "error": str(e)})
                continue
        
        return all_items
    
    # 并行处理
    all_items = []
    max_workers = min(config.PARALLEL_WORKERS, len(pending_files))
    
    logger.info(f"开始并行处理{file_type}文件", {"file_count": len(pending_files), "workers": max_workers})
    
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        # 提交所有任务
        future_to_file = {executor.submit(tracked_processor, file): file for file in pending_files}
        
        # 收集结果（带进度指示）
        completed_count = 0
        total_files = len(pending_files)
        
        # 每处理10%或至少每5个文件记录一次进度
        progress_interval = max(5, total_files // 10) if total_files > 10 else 1
        
        for future in concurrent.futures.as_completed(future_to_file):
            file = future_to_file[future]
            try:
                items = future.result()
                all_items.extend(items)
                logger.debug(f"{file_type}文件处理成功", {"file": file.name, "items": len(items)})
            except Exception as e:
                logger.error(f"{file_type}文件并行处理失败", {"file": file.name, "error": str(e)})
            
            # 进度更新
            completed_count += 1
            if completed_count % progress_interval == 0 or completed_count == total_files:
                progress_percent = (completed_count / total_files) * 100
                overall = progress_tracker.get_overall_progress()
                logger.info(f"{file_type}文件处理进度", {
                    "completed": completed_count,
                    "total": total_files,
                    "progress": f"{progress_percent:.1f}%",
                    "overall_files": f"{overall['completed_files']}/{overall['total_files']}",
                    "overall_questions": f"{overall['completed_questions']}/{overall['total_questions']}"
                })
    
    return all_items

# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------

def load_raw_questions() -> List[Dict[str, Any]]:
    """加载所有题目并返回结构化数据"""
    all_items = []
    
    logger.info("开始提取题目", {"input_dir": str(config.INPUT_DIR)})
    
    # 并行处理PDF文件
    pdf_files = list(config.INPUT_DIR.glob("*.pdf"))
    if pdf_files:
        pdf_items = process_files_parallel(pdf_files, process_single_pdf, "PDF")
        all_items.extend(pdf_items)
        logger.info("PDF文件处理完成", {"file_count": len(pdf_files), "question_count": len(pdf_items)})
    
    # 并行处理DOCX文件
    docx_files = list(config.INPUT_DIR.glob("*.docx"))
    if docx_files:
        docx_items = process_files_parallel(docx_files, process_single_docx, "DOCX")
        all_items.extend(docx_items)
        logger.info("DOCX文件处理完成", {"file_count": len(docx_files), "question_count": len(docx_items)})
    
    # 并行处理TXT文件
    txt_files = list(config.INPUT_DIR.glob("*.txt"))
    if txt_files:
        txt_items = process_files_parallel(txt_files, process_single_txt, "TXT")
        all_items.extend(txt_items)
        logger.info("TXT文件处理完成", {"file_count": len(txt_files), "question_count": len(txt_items)})
    
    logger.info("题目提取完成", {"total_questions": len(all_items)})
    return all_items

def main():
    """主入口函数"""
    try:
        logger.info("纯本地文档处理器启动")
        logger.info(f"输入目录: {config.INPUT_DIR}")
        logger.info(f"输出目录: {config.OUTPUT_DIR}")
        
        # 检查输入目录
        if not config.INPUT_DIR.exists():
            logger.error("输入目录不存在", {"path": str(config.INPUT_DIR)})
            sys.exit(1)
        
        # 显示断点续传进度
        if config.RESUME_ENABLE:
            overall_progress = progress_tracker.get_overall_progress()
            logger.info("断点续传进度", overall_progress)
        
        # 提取题目
        structured_items = load_raw_questions()
        
        if not structured_items:
            logger.warning("未提取到任何题目")
        else:
            # 保存结果
            output_path = config.OUTPUT_DIR / "extracted_questions.json"
            atomic_write_json(output_path, structured_items)
            logger.info(f"成功提取 {len(structured_items)} 道题目", {"output_file": str(output_path)})
        
        # 显示最终统计
        if config.RESUME_ENABLE:
            final_stats = progress_tracker.get_overall_progress()
            logger.info("处理完成统计", final_stats)
        
        # 显示性能统计
        performance_monitor.print_summary()
        
        # 释放OCR资源
        ocr_manager.release_all()
        
        logger.info("文档提取完成")
        
    except KeyboardInterrupt:
        logger.warning("用户中断执行")
        ocr_manager.release_all()
        sys.exit(1)
    except Exception as e:
        logger.critical("程序执行失败", {"error": str(e)})
        ocr_manager.release_all()
        sys.exit(1)

def parse_args():
    """解析命令行参数"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description="纯本地文档处理器：从PDF/DOCX/TXT源文件中读取并切分原始题目",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
环境变量优先级：命令行参数 > 环境变量 > 默认值

示例用法：
  python document_extractor.py --input ./documents --output ./results --parallel 8
  python document_extractor.py --input ./documents --resume 0 --batch-size 500
  python document_extractor.py --input ./documents --log-level DEBUG

支持的配置参数：
  INPUT_DIR          输入目录路径
  OUTPUT_DIR         输出目录路径
  PARALLEL_ENABLE    是否启用并行处理 (1/0)
  PARALLEL_WORKERS   并行工作线程数
  RESUME_ENABLE      是否启用断点续传 (1/0)
  BATCH_SIZE         大文件分批处理大小
  MAX_MEMORY_PERCENT 最大内存使用百分比
  LOG_LEVEL          日志级别 (DEBUG/INFO/WARNING/ERROR)
"""
    )
    
    # 输入输出目录
    parser.add_argument("--input", "-i", type=str, 
                       help=f"输入目录路径（默认: {config.INPUT_DIR}）")
    parser.add_argument("--output", "-o", type=str,
                       help=f"输出目录路径（默认: {config.OUTPUT_DIR}）")
    parser.add_argument("--config", "-c", type=str,
                       help="配置文件路径（YAML或JSON格式）")
    
    # 并行处理配置
    parser.add_argument("--parallel-enable", type=int, choices=[0, 1], default=None,
                       help="是否启用并行处理 (1:启用, 0:禁用)")
    parser.add_argument("--parallel-workers", "--workers", type=int, default=None,
                       help=f"并行工作线程数（默认: {config.PARALLEL_WORKERS}）")
    
    # 断点续传配置
    parser.add_argument("--resume-enable", type=int, choices=[0, 1], default=None,
                       help="是否启用断点续传 (1:启用, 0:禁用)")
    
    # 内存管理配置
    parser.add_argument("--batch-size", type=int, default=None,
                       help=f"大文件分批处理大小（默认: {config.BATCH_SIZE}）")
    parser.add_argument("--max-memory-percent", type=float, default=None,
                       help=f"最大内存使用百分比（默认: {config.MAX_MEMORY_PERCENT}）")
    
    # OCR缓存配置
    parser.add_argument("--ocr-cache-enable", type=int, choices=[0, 1], default=None,
                       help="是否启用OCR结果缓存 (1:启用, 0:禁用)")
    parser.add_argument("--ocr-cache-size", type=int, default=None,
                       help=f"OCR缓存最大条目数（默认: {config.OCR_CACHE_MAX_SIZE}）")
    parser.add_argument("--ocr-smart-selection", type=int, choices=[0, 1], default=None,
                       help="是否启用智能OCR引擎选择 (1:启用, 0:禁用)")
    
    # 日志配置
    parser.add_argument("--log-level", type=str, 
                       choices=["DEBUG", "INFO", "WARNING", "ERROR"],
                       help=f"日志级别（默认: {config.LOG_LEVEL}）")
    
    # 其他功能
    parser.add_argument("--dry-run", action="store_true",
                       help="试运行模式，只检查文件和依赖，不实际处理")
    parser.add_argument("--version", "-v", action="version", 
                       version="document_extractor v1.2.0")
    
    return parser.parse_args()

def load_config_from_file(config_path: Path) -> Dict[str, Any]:
    """从YAML或JSON配置文件加载配置"""
    if not config_path.exists():
        logger.warning(f"配置文件不存在: {config_path}")
        return {}
    
    try:
        # 根据文件扩展名选择加载器
        if config_path.suffix.lower() in ('.yaml', '.yml'):
            try:
                import yaml  # type: ignore
                with open(config_path, 'r', encoding='utf-8') as f:
                    return yaml.safe_load(f) or {}
            except ImportError:
                logger.warning("未安装PyYAML库，无法加载YAML配置文件")
                return {}
        elif config_path.suffix.lower() == '.json':
            import json
            with open(config_path, 'r', encoding='utf-8') as f:
                return json.load(f) or {}
        else:
            logger.warning(f"不支持的配置文件格式: {config_path.suffix}")
            return {}
    except Exception as e:
        logger.error(f"加载配置文件失败: {config_path}", {"error": str(e)})
        return {}

def apply_args_to_config(args):
    """将命令行参数应用到配置"""
    # 加载配置文件（如果指定）
    if args.config:
        config_path = Path(args.config).expanduser().resolve()
        file_config = load_config_from_file(config_path)
        if file_config:
            logger.info(f"从配置文件加载配置: {config_path}")
            # 应用配置文件中的配置项（仅更新存在的键）
            for key, value in file_config.items():
                if hasattr(config, key):
                    # 类型转换（保持与现有类型一致）
                    current_value = getattr(config, key)
                    try:
                        if isinstance(current_value, bool):
                            # 布尔值可能以字符串形式存储
                            if isinstance(value, str):
                                value = value.lower() in ('true', '1', 'yes', 'on')
                            setattr(config, key, bool(value))
                        elif isinstance(current_value, int):
                            setattr(config, key, int(value))
                        elif isinstance(current_value, float):
                            setattr(config, key, float(value))
                        elif isinstance(current_value, Path):
                            setattr(config, key, Path(str(value)).expanduser().resolve())
                        else:
                            setattr(config, key, value)
                    except Exception as e:
                        logger.warning(f"配置项类型转换失败: {key}={value}", {"error": str(e)})
                else:
                    logger.warning(f"忽略未知配置项: {key}")
    
    # 命令行参数覆盖配置文件和环境变量
    if args.input:
        config.INPUT_DIR = Path(args.input).expanduser().resolve()
    if args.output:
        config.OUTPUT_DIR = Path(args.output).expanduser().resolve()
    
    if args.parallel_enable is not None:
        config.PARALLEL_ENABLE = bool(args.parallel_enable)
    if args.parallel_workers is not None:
        config.PARALLEL_WORKERS = args.parallel_workers
    
    if args.resume_enable is not None:
        config.RESUME_ENABLE = bool(args.resume_enable)
    
    if args.batch_size is not None:
        config.BATCH_SIZE = args.batch_size
    if args.max_memory_percent is not None:
        config.MAX_MEMORY_PERCENT = args.max_memory_percent
    
    # OCR缓存配置
    if args.ocr_cache_enable is not None:
        config.OCR_CACHE_ENABLE = bool(args.ocr_cache_enable)
    if args.ocr_cache_size is not None:
        config.OCR_CACHE_MAX_SIZE = args.ocr_cache_size
    
    # 智能OCR引擎选择
    if args.ocr_smart_selection is not None:
        config.OCR_SMART_SELECTION = bool(args.ocr_smart_selection)
    
    if args.log_level:
        config.LOG_LEVEL = args.log_level
    
    # 确保输出目录存在
    config.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def main():
    """主入口函数"""
    try:
        # 解析命令行参数
        args = parse_args()
        apply_args_to_config(args)
        
        logger.info("纯本地文档处理器启动")
        logger.info(f"命令行配置已应用: {vars(args)}")
        logger.info(f"输入目录: {config.INPUT_DIR}")
        logger.info(f"输出目录: {config.OUTPUT_DIR}")
        
        # 检查输入目录
        if not config.INPUT_DIR.exists():
            logger.error("输入目录不存在", {"path": str(config.INPUT_DIR)})
            sys.exit(1)
        
        # 显示断点续传进度
        if config.RESUME_ENABLE:
            overall_progress = progress_tracker.get_overall_progress()
            logger.info("断点续传进度", overall_progress)
        
        # 干运行模式
        if hasattr(args, 'dry_run') and args.dry_run:
            logger.info("干运行模式：检查文件和依赖...")
            deps = check_dependencies()
            logger.info("依赖检查完成", deps)
            # 列出文件但不处理
            files = find_files(config.INPUT_DIR)
            logger.info(f"找到 {len(files)} 个文件（不处理）", {"files": [f.name for f in files]})
            sys.exit(0)
        
        # 提取题目
        structured_items = load_raw_questions()
        
        if not structured_items:
            logger.warning("未提取到任何题目")
        else:
            # 保存结果
            output_path = config.OUTPUT_DIR / "extracted_questions.json"
            atomic_write_json(output_path, structured_items)
            logger.info(f"成功提取 {len(structured_items)} 道题目", {"output_file": str(output_path)})
        
        # 显示最终统计
        if config.RESUME_ENABLE:
            final_stats = progress_tracker.get_overall_progress()
            logger.info("处理完成统计", final_stats)
        
        # 释放OCR资源
        ocr_manager.release_all()
        
        logger.info("文档提取完成")
        
    except KeyboardInterrupt:
        logger.warning("用户中断执行")
        ocr_manager.release_all()
        sys.exit(1)
    except Exception as e:
        logger.critical("程序执行失败", {"error": str(e)})
        ocr_manager.release_all()
        sys.exit(1)

if __name__ == "__main__":
    main()